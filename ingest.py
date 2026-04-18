from __future__ import annotations

import argparse
import logging
from pathlib import Path
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument

try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        TextLoader,
        UnstructuredWordDocumentLoader,
    )
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
except ImportError:  # pragma: no cover
    from langchain.document_loaders import (  # type: ignore
        PyPDFLoader,
        TextLoader,
        UnstructuredWordDocumentLoader,
    )
    from langchain.embeddings import HuggingFaceEmbeddings  # type: ignore
    from langchain.vectorstores import FAISS  # type: ignore

from langchain_openai import OpenAIEmbeddings

from utils import (
    deduplicate_documents,
    detect_file_type,
    ensure_directory,
    infer_category,
    list_supported_files,
    load_env_file,
)


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
LOGGER = logging.getLogger("placement-rag-ingest")


def build_embeddings(
    provider: str = "huggingface",
    model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
):
    """Create embeddings with a simple switch between OpenAI and Hugging Face."""
    provider = provider.lower()
    if provider == "openai":
        return OpenAIEmbeddings(model=model_name)
    if provider == "huggingface":
        return HuggingFaceEmbeddings(model_name=model_name)
    raise ValueError("Unsupported embedding provider. Use 'openai' or 'huggingface'.")


def load_docx_with_fallback(file_path: Path) -> List[Document]:
    try:
        loader = UnstructuredWordDocumentLoader(str(file_path))
        return loader.load()
    except Exception as exc:  # pragma: no cover
        LOGGER.warning(
            "Unstructured loader failed for %s; falling back to python-docx: %s",
            file_path.name,
            exc,
        )
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
            if not paragraphs:
                return []
            return [Document(page_content="\n".join(paragraphs), metadata={})]
        except Exception as fallback_exc:  # pragma: no cover
            LOGGER.exception("Fallback DOCX parsing failed for %s: %s", file_path.name, fallback_exc)
            return []


def load_single_file(file_path: Path) -> List[Document]:
    file_type = detect_file_type(file_path)
    if file_type == ".pdf":
        loader = PyPDFLoader(str(file_path))
    elif file_type == ".docx":
        documents = load_docx_with_fallback(file_path)
        loader = None
    elif file_type == ".txt":
        loader = TextLoader(str(file_path), encoding="utf-8")
    else:
        LOGGER.warning("Skipping unsupported file: %s", file_path.name)
        return []

    try:
        if loader is not None:
            documents = loader.load()
    except Exception as exc:  # pragma: no cover
        LOGGER.exception("Failed to parse %s: %s", file_path.name, exc)
        return []

    cleaned_documents: List[Document] = []
    category = infer_category(file_path)
    for document in documents:
        content = document.page_content.strip()
        if not content:
            continue
        metadata = dict(document.metadata or {})
        metadata["source"] = file_path.name
        metadata["category"] = category
        metadata["file_path"] = str(file_path.resolve())
        cleaned_documents.append(Document(page_content=content, metadata=metadata))

    if not cleaned_documents:
        LOGGER.warning("No usable text found in %s", file_path.name)
    return cleaned_documents


def load_documents(data_dir: Path) -> List[Document]:
    files = list_supported_files(data_dir)
    if not files:
        raise FileNotFoundError(
            f"No supported documents found in {data_dir}. "
            "Expected PDF, DOCX, or TXT files."
        )

    all_documents: List[Document] = []
    for file_path in files:
        LOGGER.info("Loading %s", file_path.name)
        all_documents.extend(load_single_file(file_path))

    deduped = deduplicate_documents(all_documents)
    if not deduped:
        raise ValueError("All documents were empty or failed to load.")
    return deduped


def split_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    if not chunks:
        raise ValueError("Document splitting produced zero chunks.")
    return chunks


def build_vectorstore(
    data_dir: Path,
    vectorstore_dir: Path,
    embedding_provider: str = "huggingface",
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
) -> int:
    ensure_directory(vectorstore_dir)
    documents = load_documents(data_dir)
    chunks = split_documents(documents)
    embeddings = build_embeddings(provider=embedding_provider, model_name=embedding_model)
    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(str(vectorstore_dir))
    LOGGER.info(
        "Saved vectorstore to %s with %s documents and %s chunks.",
        vectorstore_dir,
        len(documents),
        len(chunks),
    )
    return len(chunks)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Ingest placement documents into FAISS.")
    parser.add_argument(
        "--data-dir",
        default="data",
        help="Directory containing official placement documents.",
    )
    parser.add_argument(
        "--vectorstore-dir",
        default="vectorstore",
        help="Directory where the FAISS index will be stored.",
    )
    parser.add_argument(
        "--embedding-provider",
        default="huggingface",
        choices=["openai", "huggingface"],
        help="Embedding provider used to build the vectorstore.",
    )
    parser.add_argument(
        "--embedding-model",
        default="sentence-transformers/all-MiniLM-L6-v2",
        help="Embedding model name for the selected provider.",
    )
    return parser.parse_args()


def main() -> None:
    load_env_file(Path(".env"))
    args = parse_args()
    chunk_count = build_vectorstore(
        data_dir=Path(args.data_dir),
        vectorstore_dir=Path(args.vectorstore_dir),
        embedding_provider=args.embedding_provider,
        embedding_model=args.embedding_model,
    )
    print(f"Ingestion complete. Indexed {chunk_count} chunks.")


if __name__ == "__main__":
    main()
